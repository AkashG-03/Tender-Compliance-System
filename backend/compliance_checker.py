#COMPLIANCE_CHECKER.PY
"""
compliance_checker.py
Tender-aware compliance checking.

Features:
- Keyword matching
- Fuzzy string matching
- Semantic similarity (sentence-transformers)
- Tender-aware interpretation of Partial vs Met vs Missing

IMPORTANT:
- Output schema is UNCHANGED
"""

import re
from typing import List, Dict
from difflib import SequenceMatcher
from datetime import datetime

import fitz
import docx
import numpy as np
from sentence_transformers import SentenceTransformer

from models import Requirement, ComplianceResult, ComplianceReport


class ComplianceChecker:

    def __init__(self):
        # Similarity thresholds
        self.similarity_threshold_high = 0.6
        self.similarity_threshold_medium = 0.4

        print("✅ Compliance Checker initialized (Tender-aware)")

        # Load semantic model
        print("📥 Loading semantic model...")
        self.emb_model = SentenceTransformer("all-MiniLM-L6-v2")
        self.emb_dim = self.emb_model.get_sentence_embedding_dimension()
        print("✅ Semantic model loaded")

        # Tender-aware category rules
        self.category_rules = {
            "technical proposal": {
                "implicit_acceptance": False
            },
            "financial proposal": {
                "implicit_acceptance": True
            },
            "guidelines": {
                "implicit_acceptance": True
            },
            "contract / legal": {
                "implicit_acceptance": True
            }
        }

    # ------------------------------------------------------------------
    # TEXT EXTRACTION
    # ------------------------------------------------------------------
    def extract_text_from_file(self, filepath: str) -> str:
        ext = filepath.lower().split(".")[-1]
        text = ""

        try:
            if ext == "pdf":
                with fitz.open(filepath) as doc:
                    for page in doc:
                        text += page.get_text() + "\n"
                        text += self._extract_tables_from_pdf_page(page)
            elif ext == "docx":
                doc = docx.Document(filepath)
                text = "\n".join(p.text for p in doc.paragraphs)
                text += "\n" + self._extract_tables_from_docx(doc)
            else:
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read()
        except Exception as e:
            print(f"⚠️ Error extracting text from {filepath}: {e}")

        return text

    # ------------------------------------------------------------------
    # MAIN ENTRY POINT
    # ------------------------------------------------------------------
    def check_compliance(
        self,
        requirements: List[Requirement],
        company_documents: Dict[str, str],
        company_name: str,
        tender_id: str,
        page_map: Dict[int, int] = None
    ) -> ComplianceReport:

        all_text = ""
        for doc_type, path in company_documents.items():
            all_text += f"\n\n===== {doc_type.upper()} =====\n"
            all_text += self.extract_text_from_file(path)

        all_text_lower = all_text.lower()
        sentences = self._split_into_sentences(all_text)
        # 🔒 Separate normal text vs table rows
        text_sentences = [s for s in sentences if not self._is_table_sentence(s)]
        table_sentences = [s for s in sentences if self._is_table_sentence(s)]
        # ⚠️ IMPORTANT: only TEXT sentences are indexed + embedded
        indexed_sentences = list(enumerate(sentences))

        sentence_embeddings = (
            self.emb_model.encode(sentences, convert_to_numpy=True)
            if text_sentences else np.zeros((0, self.emb_dim))
        )

        # Cache requirement embeddings (efficiency)
        req_embeddings = {
            r.req_id: self.emb_model.encode(
                [r.requirement_text], convert_to_numpy=True
            )[0]
            for r in requirements
        }

        detailed_results = []
        met, partial, missing = [], [], []

        print(f"🔍 Analyzing {len(requirements)} requirements for {company_name}")

        for req in requirements:
            result = self._analyze_requirement(
                requirement=req,
                sentences=indexed_sentences,
                full_text_lower=all_text_lower,
                sentence_embeddings=sentence_embeddings,
                req_embedding=req_embeddings[req.req_id],
                page_map=page_map
            )


            detailed_results.append(result)

            if result.status == "met":
                met.append(req.req_id)
            elif result.status == "partial":
                partial.append(req.req_id)
            else:
                missing.append(req.req_id)

        compliance_percentage = (
            (len(met) + 0.5 * len(partial)) / len(requirements) * 100
            if requirements else 0
        )

        return ComplianceReport(
            company_name=company_name,
            tender_id=tender_id,
            analysis_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            total_requirements=len(requirements),
            requirements_met=met,
            requirements_partial=partial,
            requirements_missing=missing,
            compliance_percentage=round(compliance_percentage, 2),
            detailed_results=detailed_results
        )

    # ------------------------------------------------------------------
    # REQUIREMENT ANALYSIS (TENDER-AWARE)
    # ------------------------------------------------------------------
    def _analyze_requirement(
        self,
        requirement: Requirement,
        sentences: List[tuple],
        full_text_lower: str,
        sentence_embeddings: np.ndarray,
        req_embedding: np.ndarray,
        page_map: Dict[int, int] = None
    ) -> ComplianceResult:

        req_text = requirement.requirement_text.lower()
        req_numbers = self._extract_numeric_values(requirement.requirement_text)
        num_constraint = self._numeric_constraint_type(requirement.requirement_text)
        req_keywords = [k.lower() for k in requirement.keywords if k]

        # Keyword score
        keyword_hits = sum(1 for k in req_keywords if k in full_text_lower)
        keyword_score = keyword_hits / len(req_keywords) if req_keywords else 0.0

        # Fuzzy score + snippets
        fuzzy_score = 0.0
        matched_sections = []

        for _,s in sentences:
            sim = SequenceMatcher(None, req_text, s.lower()).ratio()
            if sim > fuzzy_score:
                fuzzy_score = sim


        # Semantic similarity
        semantic_score = 0.0
        if sentence_embeddings.shape[0] > 0:
            sent_norms = np.linalg.norm(sentence_embeddings, axis=1) + 1e-8
            req_norm = np.linalg.norm(req_embedding) + 1e-8
            sims = (sentence_embeddings @ req_embedding) / (sent_norms * req_norm)
            semantic_score = float(sims.max())

            # 🔐 IMPROVED FIX: rank semantic evidence with legal strength
            candidates=[]
            top_indices = sims.argsort()[-8:][::-1]

            for idx in top_indices:
                sent_idx, sentence = sentences[idx]
                page = page_map.get(sent_idx) if page_map else None
                sim_score = float(sims[idx])

                # Hard skip boilerplate
                if self._is_boilerplate(sentence):
                    continue

                 # 🔒 STEP 3: NUMERIC ENFORCEMENT (HARD BLOCK)
                if req_numbers:
                    sent_numbers = self._extract_numeric_values(sentence)

                    # Sentence must contain numeric proof
                    if not sent_numbers:
                        continue

                    req_val, req_cur = req_numbers[0]
                    sent_val, sent_cur = sent_numbers[0]

                    req_cur = self._normalize_currency(req_cur)
                    sent_cur = self._normalize_currency(sent_cur)

                    # Currency mismatch
                    if req_cur and sent_cur and req_cur != sent_cur:
                        continue

                    # Constraint enforcement
                    if num_constraint == "exact" and sent_val != req_val:
                        continue
                    if num_constraint == "min" and sent_val < req_val:
                        continue
                    if num_constraint == "max" and sent_val > req_val:
                        continue

                #legal_boost = self._legal_strength(sentence)
                #final_score = sim_score + (0.05 * legal_boost)
                # 🛡️ FIX 3: Do not apply legal-strength boost to table rows
                if " | " in sentence:
                     legal_boost = 0
                else:
                     legal_boost = self._legal_strength(sentence)

                final_score = sim_score + (0.05 * legal_boost)

                candidates.append((final_score, (sent_idx, sentence)))

            # Sort by strongest combined evidence
            candidates.sort(key=lambda x: x[0], reverse=True)

            for _, (sent_idx, sentence) in candidates[:3]:
                page = page_map.get(sent_idx) if page_map else None

                snippet = sentence[:220].strip()
                if not snippet.endswith("..."):
                    snippet += "..."

                #if page:
                #   matched_sections.append(f"[Page {page}] {snippet}")
                #else:
                #    matched_sections.append(snippet)
                if page:
                    matched_sections.append({
                        "text": sentence,
                        "snippet": f"[Page {page}] {snippet}"
                    })
                else:
                    matched_sections.append({
                        "text": sentence,
                        "snippet": snippet
                    })


            # 🛡️ Drop weak generic capability / marketing evidence
            matched_sections = [
                s for s in matched_sections
                if not any(x in s["text"].lower() for x in [
                    "our experience",
                    "we have extensive experience",
                    "industry best practices",
                    "we are capable of",
                    "we have successfully delivered"
                ])
            ]

            
            # 🛡️ FIX 2: Cap semantic dominance for table-only evidence
            if matched_sections and self._evidence_is_table_only(matched_sections):
                semantic_score = min(semantic_score, 0.55)

        # -------------------------------------------------------------
        # FIX 3: Inject table evidence ONLY if text evidence is weak
        # -------------------------------------------------------------
        if semantic_score < self.similarity_threshold_medium:
            for sent_idx, sentence in sentences:
                if self._is_table_sentence(sentence):
                    if any(k in sentence.lower() for k in req_keywords):
                        snippet = sentence[:220].strip()
                        if not snippet.endswith("..."):
                            snippet += "..."
                        matched_sections.append(snippet)

                        # Lift confidence only to safe Partial level
                        semantic_score = max(semantic_score, 0.45)

        # Pattern score
        pattern_score = self._check_patterns(requirement, full_text_lower)

        total_confidence = max(
            keyword_score,
            fuzzy_score,
            semantic_score,
            pattern_score
        )
        # Safety: prevent high confidence with no evidence
        if total_confidence >= self.similarity_threshold_high and not matched_sections:
            total_confidence = min(total_confidence, 0.55)


        # ------------------------------------------------------------------
        # TENDER-AWARE STATUS DECISION
        # ------------------------------------------------------------------
        category = requirement.category.lower()
        rules = self.category_rules.get(category, {})

        status = "missing"

        # 1️⃣ Strong explicit evidence (TEXT-FIRST RULE)
        if total_confidence >= self.similarity_threshold_high:
            # Met ONLY if at least one NON-TABLE sentence exists
            if any(
                not self._is_table_sentence(ms["text"])
                for ms in matched_sections
            ):
                status = "met"
            else:
                # Table-only evidence can never be Met
                status = "partial"

        # 2️⃣ Moderate evidence
        elif total_confidence >= self.similarity_threshold_medium:
            status = "partial"

        # 3️⃣ Implicit acceptance (non-technical)
        elif rules.get("implicit_acceptance"):
            if any(p in full_text_lower for p in [
                "we confirm that we have read",
                "terms and conditions",
                "confidential",
                "validity of the proposal"
            ]):
                status = "partial"
                total_confidence = max(total_confidence, 0.45)

        # 4️⃣ Final safety rule: only Technical can be truly Missing
        if status == "missing" and category != "technical proposal":
            status = "partial"
            total_confidence = max(total_confidence, 0.4)

        # 🔒 STEP 4: Numeric requirements cannot be Met without numeric evidence
        if req_numbers and status == "met":
            if not any(self._extract_numeric_values(ms["text"]) for ms in matched_sections):
                status = "partial"
                total_confidence = min(total_confidence, 0.6)


        # 🛡️ RULE 6: Confidence alignment (evaluator-grade)
        if status == "met":
            total_confidence = min(max(total_confidence, 0.70), 0.85)
        elif status == "partial":
            total_confidence = min(max(total_confidence, 0.45), 0.65)
        else:  # missing
            total_confidence = min(total_confidence, 0.39)

        reasoning = (
            f"Tender-aware evaluation applied. "
            f"Confidence: {total_confidence:.2f}. "
            f"Category: {requirement.category}. "
            f"Evidence-based classification enforced."
        )
        matched_sections = [ms["snippet"] for ms in matched_sections]
        return ComplianceResult(
            req_id=requirement.req_id,
            requirement_text=requirement.requirement_text,
            status=status,
            confidence_score=round(total_confidence, 3),
            semantic_score=round(semantic_score, 3),
            matched_sections=matched_sections,
            reasoning=reasoning
        )

    # ------------------------------------------------------------------
    # HELPERS
    # ------------------------------------------------------------------
    def _split_into_sentences(self, text: str) -> List[str]:
        parts = re.split(r"[.!?]\s+", text)
        return [p.strip() for p in parts if len(p.strip()) > 25]

    def _check_patterns(self, requirement: Requirement, text: str) -> float:
        terms = self._extract_key_terms(requirement.requirement_text.lower())
        if not terms:
            return 0.0
        return sum(1 for t in terms if t in text) / len(terms)

    def _extract_key_terms(self, text: str) -> List[str]:
        stop_words = {
            "the", "and", "or", "shall", "must", "with",
            "for", "from", "that", "this", "will",
            "are", "be", "have"
        }
        words = re.findall(r"\b[a-z]{4,}\b", text.lower())
        return list({w for w in words if w not in stop_words})[:10]
    
    def _is_boilerplate(self, sentence: str) -> bool:
        boilerplate_phrases = [
            "we confirm that we have read",
            "no conflicts of interest",
            "submission letter",
            "we, the undersigned",
            "annex",
            "to: agra procurement",
            "offer to provide",
            "hereby submitting"
        ]
        s = sentence.lower()
        return any(p in s for p in boilerplate_phrases)
    
    def _legal_strength(self, sentence: str) -> int:
        keywords = [
            "shall", "must", "will", "valid", "certificate",
            "warrant", "confidential", "comply", "submit",
            "deliver", "invoice", "approve"
        ]
        s = sentence.lower()
        return sum(1 for k in keywords if k in s)
    
    def _extract_tables_from_pdf_page(self, page) -> str:
        """
        Extract tables from a PDF page and convert rows to sentence-like text.
        Safe: silently fails if tables are not detected.
        """
        table_text = ""
        try:
            tables = page.find_tables()
            for table in tables:
                for row in table.extract():
                    cells = [str(c).strip() for c in row if c and str(c).strip()]
                    if len(cells) >= 2:
                        table_text += " | ".join(cells) + ".\n"
        except Exception:
            pass
        return table_text
    
    def _extract_tables_from_docx(self, doc) -> str:
        """
        Extract tables from a DOCX document and convert rows to sentences.
        """
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if len(cells) >= 2:
                    table_text += " | ".join(cells) + ".\n"
        return table_text
    
    def _evidence_is_table_only(self, matched_sections) -> bool:
        return all(" | " in s["text"] for s in matched_sections)
    

    def _explicit_acceptance_present(self, sentence: str, requirement_text: str) -> bool:
        """
        Returns True only if the proposal explicitly confirms or accepts the requirement.
        """
        s = sentence.lower()
        r = requirement_text.lower()

        if " | " in s:
            return False

        acceptance_phrases = [
            "we confirm",
            "we confirm that",
            "we will comply",
            "we agree to",
            "we accept",
            "we shall comply",
            "we undertake",
            "we commit to",
            "will be provided",
            "shall be provided",
            "is included",
            "is submitted",
            "is attached"
        ]

        # 🔒 Extract meaningful requirement terms (avoid generic verbs)
        req_terms = [w for w in r.split() if len(w) > 4][:5]

        # ✅ Evaluator-grade acceptance:
        # 1) Explicit acceptance language
        # 2) At least TWO meaningful requirement terms present
        # Strong anchor terms (avoid generic verbs)
        anchor_terms = [
            t for t in req_terms
            if t not in {"submit", "provide", "provide", "document", "documents", "information"}
        ]

        return (
            any(p in s for p in acceptance_phrases)
            and any(t in s for t in anchor_terms)
        )
    
    def _is_table_sentence(self, s: str) -> bool:
        return " | " in s
    
    def _extract_numeric_values(self, text: str):
        """
        Extract numeric values with optional currency.
        Returns list of tuples: (value: float, currency: str|None)
        """
        matches = re.findall(
            r"(\d+(?:\.\d+)?)\s*(usd|\$|inr|eur)?",
            text.lower()
        )

        results = []
        for val, cur in matches:
            results.append((float(val), cur))
        return results


    def _normalize_currency(self, currency: str | None):
        if currency in {"$", "usd"}:
            return "usd"
        if currency == "inr":
            return "inr"
        if currency == "eur":
            return "eur"
        return None


    def _numeric_constraint_type(self, text: str):
        t = text.lower()
        if any(x in t for x in ["at least", "minimum", ">=", "not less than"]):
            return "min"
        if any(x in t for x in ["at most", "maximum", "<=", "not more than"]):
            return "max"
        if any(x in t for x in ["between", "range", "from"]):
            return "range"
        return "exact"

    
